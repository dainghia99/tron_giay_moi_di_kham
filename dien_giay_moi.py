# -*- coding: utf-8 -*-
"""
Quy tắc thuật toán xử lý dữ liệu:
  - Gom các dòng CSV thành từng HỘ: dòng có STT là chủ hộ, các dòng bỏ trống STT
    phía dưới là thành viên của hộ đó.
  - Chỉ tạo giấy mời cho hộ có ÍT NHẤT 1 người "Chưa khám" (kể cả chủ hộ).
  - Liệt kê TẤT CẢ thành viên, ghi đúng trạng thái từng người.
  - Chủ hộ luôn đứng đầu.
"""

import csv
import os
import re
import copy

from docx import Document

# ------------------------------------------------------------------ CẤU HÌNH
TEMPLATE_PATH = "gm.docx"        # File Word MẪU của bạn (giữ nguyên định dạng)
CSV_PATH      = "output/yen.csv"     # File dữ liệu
OUTPUT_DIR    = "giay_moi"       # Thư mục kết quả
SO_HIEU       = "1914"           # Số hiệu — GIỮ NGUYÊN cho TẤT CẢ giấy mời
CHUA_KHAM     = "Chưa khám"

# Giá trị MẪU đang có sẵn trong file gm.docx — dùng làm "mốc" để thay thế.
# Nếu file mẫu của bạn dùng tên/ngày/địa chỉ mẫu khác, chỉ cần sửa lại đây.
MAU_TEN_CHU_HO = "Nguyễn Văn A"
MAU_NGAY_SINH  = "01/01/1900"
MAU_DIA_CHI    = "Bản Nậm Pan, xã Mường Toong, tỉnh Điện Biên"
MAU_SO_HIEU    = "1914" #Nhập số hiệu ủy ban ban hành
# ---------------------------------------------------------------------------


# ---------------------------------------------------- Thay text giữ định dạng
def replace_in_paragraph(paragraph, old, new):
    """Thay 'old' -> 'new' trong 1 đoạn văn, cố gắng giữ định dạng của run."""
    if old not in paragraph.text:
        return False
    for run in paragraph.runs:            # old nằm gọn trong 1 run -> giữ nguyên format
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    new_full = paragraph.text.replace(old, new)   # old bị tách qua nhiều run
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def replace_everywhere(paragraphs, old, new):
    done = False
    for p in paragraphs:
        if replace_in_paragraph(p, old, new):
            done = True
    return done


def iter_table_cell_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p


# ------------------------------------------- Tìm bảng danh sách thành viên
def tim_bang_thanh_vien(doc):
    for t in doc.tables:
        header = " ".join(c.text for c in t.rows[0].cells).lower()
        if "họ và tên" in header and "trạng thái" in header:
            return t
    return None


def set_cell_text(cell, text):
    """Đặt text cho ô, giữ định dạng của run đầu tiên trong ô."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:      # xoá đoạn thừa nếu ô mẫu nhiều dòng
        extra._element.getparent().remove(extra._element)


def dien_bang_thanh_vien(table, members):
    """Giữ dòng tiêu đề, dùng dòng dữ liệu đầu làm mẫu, nhân bản theo thành viên."""
    data_rows = table.rows[1:]
    if not data_rows:
        raise RuntimeError("Bảng danh sách trong file mẫu không có dòng dữ liệu mẫu.")

    mau_tr = copy.deepcopy(data_rows[0]._tr)   # dòng mẫu (giữ mọi định dạng)
    tbl = table._tbl
    for r in data_rows:                        # xoá các dòng dữ liệu cũ
        tbl.remove(r._tr)

    for i, m in enumerate(members, start=1):
        tbl.append(copy.deepcopy(mau_tr))
        row = table.rows[-1]
        vals = [str(i), m["ho_ten"], m["nam_sinh"], m["dia_chi"],
                m["thanh_phan"], m["trang_thai"]]
        for cell, v in zip(row.cells, vals):
            set_cell_text(cell, v)


# ------------------------------------------------------- ĐỌC & GOM NHÓM CSV
def doc_va_gom_ho(csv_path):
    with open(csv_path, encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        cols = {c.strip(): c for c in reader.fieldnames}
        ho_list, ho = [], None
        for row in reader:
            g = lambda k: (row.get(cols.get(k, k), "") or "").strip()
            if not g("Họ và tên"):
                continue
            member = {
                "ho_ten":     g("Họ và tên"),
                "nam_sinh":   g("Ngày tháng năm sinh"),
                "dia_chi":    g("Địa chỉ"),
                "thanh_phan": g("Thành phần"),
                "trang_thai": g("Trạng thái"),
            }
            if g("STT"):
                ho = [member]; ho_list.append(ho)
            elif ho is None:
                ho = [member]; ho_list.append(ho)
            else:
                ho.append(member)
    return ho_list


def sap_xep_ho(members):
    chu_ho = [m for m in members if m["thanh_phan"].lower() == "chủ hộ"]
    khac   = [m for m in members if m["thanh_phan"].lower() != "chủ hộ"]
    return chu_ho + khac


def slugify(name):
    s = re.sub(r"\s+", "_", name.strip())
    return re.sub(r"[^0-9A-Za-zĐđ_À-ỹ]", "", s)

# ------------------------------------------------------------- TẠO 1 GIẤY MỜI
def tao_giay_moi(members, so_hieu):
    members = sap_xep_ho(members)
    chu_ho = members[0]

    doc = Document(TEMPLATE_PATH)   # mở đúng file mẫu -> giữ nguyên định dạng

    # 1) Thông tin chủ hộ ở phần nội dung (các đoạn văn ngoài bảng)
    replace_everywhere(doc.paragraphs, MAU_TEN_CHU_HO, chu_ho["ho_ten"])
    replace_everywhere(doc.paragraphs, MAU_NGAY_SINH, chu_ho["nam_sinh"])
    replace_everywhere(doc.paragraphs, MAU_DIA_CHI,   chu_ho["dia_chi"])

    bang_tv = tim_bang_thanh_vien(doc)
    if bang_tv is None:
        raise RuntimeError("Không tìm thấy bảng danh sách thành viên trong file mẫu.")

    # 2) Số hiệu ở bảng tiêu đề (mọi bảng trừ bảng danh sách) — giữ nguyên cho mọi hộ
    for t in doc.tables:
        if t is bang_tv:
            continue
        replace_everywhere(iter_table_cell_paragraphs(t), MAU_SO_HIEU, str(so_hieu))

    # 3) Điền bảng danh sách thành viên (nhân bản dòng mẫu)
    dien_bang_thanh_vien(bang_tv, members)
    return doc


# ---------------------------------------------------------------------- MAIN
def main():
    if not os.path.exists(TEMPLATE_PATH):
        raise SystemExit(f"Không thấy file mẫu '{TEMPLATE_PATH}'. "
                         f"Hãy đặt file Word gốc của bạn cùng thư mục với script.")
    ho_list = doc_va_gom_ho(CSV_PATH)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    stt = 1                       # chỉ dùng để đặt tên file cho khỏi trùng
    tao = bo_qua = 0
    for members in ho_list:
        co_chua_kham = any(m["trang_thai"].strip().lower() == CHUA_KHAM.lower()
                           for m in members)
        chu_ho = sap_xep_ho(members)[0]
        if not co_chua_kham:
            bo_qua += 1
            print(f"  (bỏ qua) Hộ {chu_ho['ho_ten']}: tất cả đã khám")
            continue

        doc = tao_giay_moi(members, SO_HIEU)   # số hiệu GIỮ NGUYÊN cho mọi hộ
        ten_file = f"GiayMoi_{stt:02d}_{slugify(chu_ho['ho_ten'])}.docx"
        doc.save(os.path.join(OUTPUT_DIR, ten_file))
        print(f"  ✓ Đã tạo: {ten_file}  (Số {SO_HIEU}, {len(members)} thành viên)")
        tao += 1
        stt += 1

    print(f"\nHoàn tất: tạo {tao} giấy mời, bỏ qua {bo_qua} hộ (đã khám hết).")
    print(f"Kết quả trong thư mục: {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()